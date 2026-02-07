import os
import hashlib
import logging
import tempfile
import zipfile
import json
import csv
import re
import time
from typing import List, Dict, Optional, Tuple
from dataclasses import dataclass
from collections import OrderedDict
import pickle
from dotenv import load_dotenv
import nltk
import docx
import whisper
import pdfplumber
import pytesseract
import cv2
import numpy as np

from nltk.tokenize import sent_tokenize
from llama_index.core import VectorStoreIndex, Document, StorageContext, load_index_from_storage, Settings
from llama_index.embeddings.huggingface import HuggingFaceEmbedding
from llama_index.llms.groq import Groq
from llama_index.llms.ollama import Ollama
from sentence_transformers import CrossEncoder

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
load_dotenv(os.path.join(BASE_DIR, ".env"))

nltk.download("punkt", quiet=True) # tokenizer model
nltk.download("punkt_tab", quiet=True)

# === Configuration ===
PERSIST_DIR = "./storage"
CACHE_DIR = "./cache"
HASH_CACHE_FILE = "text_hash.txt"
EMBEDDING_MODEL = "sentence-transformers/all-MiniLM-L6-v2"
GROQ_MODEL = "llama-3.3-70b-versatile"
GROQ_API_KEY = os.getenv("GROQ_API_KEY")
OLLAMA_MODEL = "gemma2:2b"
CHUNK_SIZE = 512
CHUNK_OVERLAP = 50

# Retrieval parameters
K_INIT = 20
K_RERANK = 10
K_FINAL = 5
SIMILARITY_THRESHOLD = 0.7
CACHE_SIZE = 1000
CACHE_SIMILARITY_THRESHOLD = 0.95

logging.basicConfig(level=logging.WARNING, format="%(asctime)s - %(levelname)s - %(message)s") # warning

# === Cache System ===
@dataclass
class CacheEntry:
    query: str
    embedding: np.ndarray
    response: str
    sources: List[str]
    timestamp: float

class MultiLevelCache:
    """L1 (memory) + L2 (disk) caching with semantic similarity"""
    
    def __init__(self, cache_dir=CACHE_DIR, max_size=CACHE_SIZE):  # Creates cache directory
        self.cache_dir = cache_dir
        self.max_size = max_size
        os.makedirs(cache_dir, exist_ok=True)
        
        self.l1_cache: OrderedDict[str, CacheEntry] = OrderedDict()  # Maintains insertion order
        self.l2_cache_file = os.path.join(cache_dir, "l2_cache.pkl")
        self._load_l2_cache() # load sefly
    
    def _load_l2_cache(self):
        if os.path.exists(self.l2_cache_file):
            try:
                with open(self.l2_cache_file, 'rb') as f:
                    self.l1_cache = pickle.load(f)
                logging.info(f"✓ Loaded {len(self.l1_cache)} entries from L2 cache")
            except Exception as e:
                logging.warning(f"Failed to load L2 cache: {e}")
    
    def _save_l2_cache(self):
        try:
            with open(self.l2_cache_file, 'wb') as f:
                pickle.dump(self.l1_cache, f)
        except Exception as e:
            logging.warning(f"Failed to save L2 cache: {e}")
    
    def _compute_similarity(self, emb1: np.ndarray, emb2: np.ndarray) -> float:
        return np.dot(emb1, emb2) / (np.linalg.norm(emb1) * np.linalg.norm(emb2)) #cosine similarity.
    
    def get(self, query: str, query_embedding: Optional[np.ndarray] = None) -> Optional[Tuple[str, List[str]]]:
        query_hash = hashlib.md5(query.encode()).hexdigest()
        
        if query_hash in self.l1_cache:
            entry = self.l1_cache[query_hash]
            self.l1_cache.move_to_end(query_hash)
            return entry.response, entry.sources
        
        if query_embedding is not None:
            for key, entry in self.l1_cache.items():
                similarity = self._compute_similarity(query_embedding, entry.embedding)
                if similarity >= CACHE_SIMILARITY_THRESHOLD:
                    return entry.response, entry.sources
        
        return None
    
    def set(self, query: str, query_embedding: np.ndarray, response: str, sources: List[str]):
        query_hash = hashlib.md5(query.encode()).hexdigest()
        
        if len(self.l1_cache) >= self.max_size:
            self.l1_cache.popitem(last=False)
        
        self.l1_cache[query_hash] = CacheEntry(
            query=query,
            embedding=query_embedding,
            response=response,
            sources=sources,
            timestamp=time.time()
        )
        
        if len(self.l1_cache) % 10 == 0:
            self._save_l2_cache()

cache = MultiLevelCache()

# Query Classification(focussed word)
class QueryClassifier:
    FACTUAL_KEYWORDS = ["what", "who", "when", "where", "define", "explain", "list", "show", "tell"]
    ANALYTICAL_KEYWORDS = ["compare", "analyze", "why", "how can", "how does", "evaluate", "difference", "relationship"]
    CONVERSATIONAL_KEYWORDS = ["hi", "hello", "thanks", "thank you", "bye", "hey"]
    
    @staticmethod
    def classify(query: str) -> Dict[str, any]:
        query_lower = query.lower().strip()
        word_count = len(query.split())
        
        is_pure_conversational = (
            word_count <= 3 and 
            any(query_lower.startswith(kw) for kw in QueryClassifier.CONVERSATIONAL_KEYWORDS)
        )
        
        intent = "factual"
        if is_pure_conversational:
            intent = "conversational"
        elif any(kw in query_lower for kw in QueryClassifier.ANALYTICAL_KEYWORDS):
            intent = "analytical"
        elif any(kw in query_lower for kw in QueryClassifier.FACTUAL_KEYWORDS):
            intent = "factual"
        
        complexity = min(1.0, (word_count / 20) + (0.3 if intent == "analytical" else 0))
        
        return {
            "intent": intent,
            "complexity": complexity,
            "word_count": word_count
        }

# === Initialize Settings ===
def initialize_settings(use_groq=True):
    try:
        Settings.embed_model = HuggingFaceEmbedding(model_name=EMBEDDING_MODEL)
        Settings.chunk_size = CHUNK_SIZE
        Settings.chunk_overlap = CHUNK_OVERLAP
        
        if use_groq:
            try:
                if not GROQ_API_KEY or GROQ_API_KEY == "your_groq_api_key_here":
                    raise ValueError("Missing GROQ_API_KEY")
                
                Settings.llm = Groq(
                    model=GROQ_MODEL,
                    api_key=GROQ_API_KEY,
                    temperature=0.3,
                    max_tokens=600
                )
                logging.info(f"Using Groq LLM: {GROQ_MODEL}")
                return "groq"
            except Exception as e:
                logging.warning(f"Groq failed: {e}. Falling back to Ollama...")
                raise
    except Exception:
        Settings.llm = Ollama(model=OLLAMA_MODEL, temperature=0.3)
        logging.info(f"Using Ollama LLM: {OLLAMA_MODEL}")
        return "ollama"

# === File Extractors ===
def extract_text_from_pdf(path):
    try:
        with pdfplumber.open(path) as pdf:
            return "\n".join([p.extract_text() or "" for p in pdf.pages])
    except Exception as e:
        logging.error(f"PDF extraction failed: {e}")
        return ""

def extract_text_from_image(path):
    try:
        img = cv2.imread(path)
        return pytesseract.image_to_string(img)
    except Exception as e:
        logging.error(f"Image OCR failed: {e}")
        return ""

def extract_text_from_docx(path):
    try:
        return "\n".join([p.text for p in docx.Document(path).paragraphs])
    except Exception as e:
        logging.error(f"DOCX extraction failed: {e}")
        return ""

def extract_text_from_txt(path):
    try:
        with open(path, "r", encoding="utf-8") as f:
            return f.read()
    except Exception as e:
        logging.error(f"TXT extraction failed: {e}")
        return ""

def extract_text_from_csv(path):
    try:
        with open(path, newline='', encoding='utf-8') as csvfile:
            return "\n".join([" | ".join(row) for row in csv.reader(csvfile)])
    except Exception as e:
        logging.error(f"CSV extraction failed: {e}")
        return ""

def extract_text_from_json(path):
    try:
        with open(path, "r", encoding="utf-8") as f:
            return json.dumps(json.load(f), indent=2)
    except Exception as e:
        logging.error(f"JSON extraction failed: {e}")
        return ""

def extract_text_from_audio(path):
    try:
        model = whisper.load_model("base")
        return model.transcribe(path)["text"]
    except Exception as e:
        logging.error(f"Audio transcription failed: {e}")
        return ""

def extract_text_from_zip(path):
    text = ""
    try:
        with zipfile.ZipFile(path, 'r') as zip_ref:
            with tempfile.TemporaryDirectory() as temp_dir:
                zip_ref.extractall(temp_dir)
                for root, _, files in os.walk(temp_dir):
                    for name in files:
                        full_path = os.path.join(root, name)
                        text += extract_text(full_path) + "\n"
    except Exception as e:
        logging.error(f"ZIP extraction failed: {e}")
    return text

def extract_text(path):
    ext = os.path.splitext(path)[1].lower()
    extractors = {
        ".pdf": extract_text_from_pdf,
        ".jpg": extract_text_from_image,
        ".jpeg": extract_text_from_image,
        ".png": extract_text_from_image,
        ".docx": extract_text_from_docx,
        ".txt": extract_text_from_txt,
        ".csv": extract_text_from_csv,
        ".json": extract_text_from_json,
        ".mp3": extract_text_from_audio,
        ".wav": extract_text_from_audio,
        ".mp4": extract_text_from_audio,
        ".zip": extract_text_from_zip
    }
    
    extractor = extractors.get(ext)
    return extractor(path) if extractor else ""

# Utility Functions 
def clean_text(text):
    return re.sub(r'\s+', ' ', text).replace("•", "-").replace("–", "-").strip()

def compute_hash(text):
    return hashlib.md5(text.encode("utf-8")).hexdigest()  #text into bytes

def has_file_changed(text):
    if os.path.exists(HASH_CACHE_FILE):
        with open(HASH_CACHE_FILE, "r") as f:
            return compute_hash(text) != f.read().strip()
    return True

def cache_hash(text):
    with open(HASH_CACHE_FILE, "w") as f:
        f.write(compute_hash(text))

# === Context Compression ===
def compress_context(chunks: List[str], max_tokens: int = 1500) -> str:
    compressed = []
    token_count = 0
    
    for chunk in chunks:
        chunk_tokens = len(chunk.split())
        if token_count + chunk_tokens > max_tokens:
            remaining = max_tokens - token_count
            compressed.append(" ".join(chunk.split()[:remaining]))
            break
        compressed.append(chunk)
        token_count += chunk_tokens
    
    unique_sentences = []
    seen = set()
    for chunk in compressed:
        for sent in sent_tokenize(chunk):
            sent_clean = sent.strip().lower()
            if sent_clean not in seen and len(sent_clean) > 20:
                seen.add(sent_clean)
                unique_sentences.append(sent)
    
    return " ".join(unique_sentences)

# === Index Creation ===
def create_index(file_paths):
    documents = []
    for path in file_paths:
        text = clean_text(extract_text(path))
        if text:
            doc = Document(text=text, metadata={"source": os.path.basename(path)})
            documents.append(doc)
    
    if not documents:
        logging.warning("No valid text extracted from files.")
        return None
    
    combined_text = "\n".join([doc.text for doc in documents])
    
    if os.path.exists(PERSIST_DIR) and not has_file_changed(combined_text):
        logging.info("Loading existing index...")
        storage_context = StorageContext.from_defaults(persist_dir=PERSIST_DIR)
        index = load_index_from_storage(storage_context)
    else:
        logging.info("Creating new index...")
        index = VectorStoreIndex.from_documents(documents, show_progress=True)
        index.storage_context.persist(persist_dir=PERSIST_DIR)
        cache_hash(combined_text)
    
    return index

# === Response Formatter ===
def format_response(response: str, sources: List[str]) -> str:
    """Format response naturally without ChatGPT-like structure"""
    response = response.strip()
    
    # Remove structured formatting artifacts
    response = re.sub(r'^(Answer|Response|Here\'s|Based on):\s*', '', response, flags=re.IGNORECASE)
    response = re.sub(r'\*\*([^*]+)\*\*', r'\1', response)  # Remove bold markdown
    response = re.sub(r'__([^_]+)__', r'\1', response)  # Remove underline markdown
    
    # Clean up excessive formatting
    response = re.sub(r'\n{3,}', '\n\n', response)  # Max 2 consecutive newlines
    
    # Add sources inline if available
    if sources:
        unique_sources = list(dict.fromkeys(sources))
        source_text = f" [Source: {', '.join(unique_sources)}]"
        response += source_text
    
    return response

# === Optimized Query Pipeline ===
def query_index_optimized(index, query: str, llm_type: str) -> Tuple[str, List[str], Dict[str, float]]:
    timings = {}
    start_total = time.time()
    
    # Query Classification
    t0 = time.time()
    classification = QueryClassifier.classify(query)
    timings['classification'] = time.time() - t0
    
    # Cache Lookup
    t0 = time.time()
    query_embedding = Settings.embed_model.get_text_embedding(query)
    cached = cache.get(query, np.array(query_embedding))
    timings['cache_lookup'] = time.time() - t0
    
    if cached:
        response, sources = cached
        timings['total'] = time.time() - start_total
        return response, sources, timings
    
    # Conversational handling
    if classification['intent'] == 'conversational':
        response = "Hello! Ask me anything about your documents."
        cache.set(query, np.array(query_embedding), response, [])
        timings['total'] = time.time() - start_total
        return response, [], timings
    
    # Retrieval
    t0 = time.time()
    k_retrieve = K_INIT if classification['complexity'] > 0.5 else max(K_FINAL, 8)
    
    retriever = index.as_retriever(similarity_top_k=k_retrieve)
    retrieved_nodes = retriever.retrieve(query)
    timings['retrieval'] = time.time() - t0
    
    if not retrieved_nodes:
        response = "I couldn't find relevant information in the documents to answer that."
        timings['total'] = time.time() - start_total
        return response, [], timings
    
    # Reranking
    t0 = time.time()
    top_score = retrieved_nodes[0].score if hasattr(retrieved_nodes[0], 'score') else 0
    
    if top_score < SIMILARITY_THRESHOLD and len(retrieved_nodes) > K_FINAL:
        try:
            reranker = CrossEncoder('cross-encoder/ms-marco-MiniLM-L-6-v2')
            pairs = [[query, node.text] for node in retrieved_nodes[:K_RERANK]]
            scores = reranker.predict(pairs)
            
            scored_nodes = [(node, score) for node, score in zip(retrieved_nodes[:K_RERANK], scores)]
            scored_nodes.sort(key=lambda x: x[1], reverse=True)
            final_nodes = [node for node, _ in scored_nodes[:K_FINAL]]
        except Exception as e:
            final_nodes = retrieved_nodes[:K_FINAL]
    else:
        final_nodes = retrieved_nodes[:K_FINAL]
    
    timings['reranking'] = time.time() - t0
    
    # Context Compression
    t0 = time.time()
    contexts = [node.text for node in final_nodes]
    compressed_context = compress_context(contexts, max_tokens=1200)
    timings['compression'] = time.time() - t0
    
    # Generation with strict document-only prompt
    t0 = time.time()
    prompt = (
        "You are an AI assistant. Use only the given context to answer.\n"
        "If answer is not in context, reply: 'Not available in the documentation.'\n\n"
        f"Context:\n{compressed_context}\n\n"
        f"Question: {query}\n"
        "Answer:"
    )
    
    response_obj = Settings.llm.complete(prompt)
    response = response_obj.text.strip()
    
    # Post-processing: Check if response is too generic (likely hallucinated)
    generic_phrases = [
        "in general", "typically", "usually", "commonly", "it is known that",
        "experts suggest", "research shows", "studies indicate"
    ]
    
    response_lower = response.lower()
    if any(phrase in response_lower for phrase in generic_phrases) and len(response.split()) < 30:
        response = "This information is not available in the provided documents."
    
    timings['generation'] = time.time() - t0
    
    # Extract sources
    sources = [node.metadata.get('source', 'Unknown') for node in final_nodes[:3]]
    
    # Cache the result
    cache.set(query, np.array(query_embedding), response, sources)
    
    timings['total'] = time.time() - start_total
    
    return response, sources, timings

# === Main Function ===
def main():
    
    try:
        llm_type = initialize_settings(use_groq=True)
    except:
        llm_type = initialize_settings(use_groq=False)
    
    file_input = input("Enter file paths (comma-separated): ").strip()
    file_paths = [f.strip() for f in file_input.split(",") if os.path.exists(f.strip())]
    
    if not file_paths:
        print("No valid file paths provided.")
        return
    
    print(f"\nProcessing {len(file_paths)} file(s)...")
    index = create_index(file_paths)
    
    if index is None:
        print("Failed to create index.")
        return
    
    while True:
        question = input("Q: ").strip()
        
        if question.lower() in ["exit", "quit", "q"]:
            print("\nGoodbye!")
            break
        
        if not question:
            continue
        
        response, sources, timings = query_index_optimized(index, question, llm_type)
        
        formatted_response = format_response(response, sources)
        print(f"\n{formatted_response}\n")
        print(f"({timings['total']*1000:.0f}ms)\n")

if __name__ == "__main__":
    main()