from flask import Flask, render_template, request, jsonify, send_from_directory, redirect, url_for, flash
import os
import sys
import time
import traceback
import logging
import re
import hashlib
import json
import pickle
import csv
import xml.etree.ElementTree as ET
from typing import List, Dict, Optional, Tuple
from dataclasses import dataclass
from collections import OrderedDict
from dotenv import load_dotenv
from werkzeug.utils import secure_filename
from flask_cors import CORS
import smtplib
from email.mime.text import MIMEText
from email.mime.multipart import MIMEMultipart
import requests
import numpy as np
import prompt
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s'
)
BASE_DIR = os.path.dirname(os.path.abspath(__file__))
try:
    import nltk
    import docx
    import pdfplumber
    import pytesseract
    import cv2
    from PIL import Image
    from nltk.tokenize import sent_tokenize
    from llama_index.core import VectorStoreIndex, Document, StorageContext, load_index_from_storage, Settings
    from llama_index.embeddings.huggingface import HuggingFaceEmbedding
    from llama_index.llms.groq import Groq as GroqLLM
    from llama_index.llms.ollama import Ollama as OllamaLLM
    from sentence_transformers import CrossEncoder
    
    nltk.download("punkt", quiet=True)
    nltk.download("punkt_tab", quiet=True)
    RAG_AVAILABLE = True
    logging.info("RAG dependencies loaded")
except ImportError as e:
    RAG_AVAILABLE = False
    logging.warning(f"RAG dependencies not available: {e}")

try:
    from groq import Groq
    GROQ_AVAILABLE = True
except ImportError:
    GROQ_AVAILABLE = False
    logging.warning("Groq not installed")

# Load environment
load_dotenv()

BASE_DIR = os.path.abspath(os.path.dirname(__file__))
UPLOAD_FOLDER = os.path.join(BASE_DIR, 'uploads')
CONVERSATION_FOLDER = os.path.join(BASE_DIR, 'conversations')
RAG_STORAGE_DIR = os.path.join(BASE_DIR, 'rag_storage')
RAG_CACHE_DIR = os.path.join(BASE_DIR, 'rag_cache')

# Enhanced file support
ALLOWED_EXTENSIONS = {
    'pdf', 'txt', 'docx', 'doc',  # Documents
    'png', 'jpg', 'jpeg', 'bmp', 'gif', 'tiff',  # Images
    'csv', 'json', 'xml',  # Data files
    'html', 'htm', 'md'  # Web/Markdown
}

# RAG Configuration
EMBEDDING_MODEL = "sentence-transformers/all-MiniLM-L6-v2"
CHUNK_SIZE = 512
CHUNK_OVERLAP = 50
K_INIT = 20
K_RERANK = 10
K_FINAL = 5
SIMILARITY_THRESHOLD = 0.7
CACHE_SIZE = 1000
CACHE_SIMILARITY_THRESHOLD = 0.95

app = Flask(
    __name__,
    static_url_path='/static',
    static_folder=os.path.join(BASE_DIR, 'static'),
    template_folder=os.path.join(BASE_DIR, 'templates')
)
CORS(app)
app.secret_key = os.getenv("FLASK_SECRET_KEY", "dev_secret_key_change_in_production")
app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER
app.config['CONVERSATION_FOLDER'] = CONVERSATION_FOLDER
app.config['MAX_CONTENT_LENGTH'] = 150 * 1024 * 1024  # 150 MB

# Create necessary directories
for directory in [UPLOAD_FOLDER, CONVERSATION_FOLDER, RAG_STORAGE_DIR, RAG_CACHE_DIR]:
    os.makedirs(directory, exist_ok=True)

GROQ_API_KEY = os.getenv("GROQ_API_KEY")
OLLAMA_API_URL = "http://localhost:11434/api/generate"

if GROQ_API_KEY:
    GROQ_API_KEY = GROQ_API_KEY.strip()
    os.environ["GROQ_API_KEY"] = GROQ_API_KEY
else:
    logging.warning("GROQ_API_KEY not set")

# Initialize Groq client
groq_client = None
if GROQ_AVAILABLE and GROQ_API_KEY and GROQ_API_KEY != "your_groq_api_key_here":
    try:
        groq_client = Groq(api_key=GROQ_API_KEY)
        logging.info("Groq client initialized")
    except Exception as e:
        logging.error(f"Groq initialization failed: {e}")

# Email Configuration
SENDER_EMAIL = os.getenv("SENDER_EMAIL")
SENDER_PASSWORD = os.getenv("SENDER_PASSWORD")
RECEIVER_EMAIL = os.getenv("RECEIVER_EMAIL", SENDER_EMAIL)

@dataclass
class CacheEntry:
    """Cache entry for storing query results"""
    query: str
    embedding: np.ndarray
    response: str
    sources: List[str]
    timestamp: float


class MultiLevelCache:
    """L1 (memory) + L2 (disk) caching with semantic similarity"""
    
    def __init__(self, cache_dir=RAG_CACHE_DIR, max_size=CACHE_SIZE):
        self.cache_dir = cache_dir
        self.max_size = max_size
        os.makedirs(cache_dir, exist_ok=True)
        
        self.l1_cache: OrderedDict[str, CacheEntry] = OrderedDict()
        self.l2_cache_file = os.path.join(cache_dir, "l2_cache.pkl")
        self._load_l2_cache()
    
    def _load_l2_cache(self):
        if os.path.exists(self.l2_cache_file):
            try:
                with open(self.l2_cache_file, 'rb') as f:
                    self.l1_cache = pickle.load(f)
                logging.info(f"Loaded {len(self.l1_cache)} cache entries")
            except Exception as e:
                logging.warning(f"Cache load failed: {e}")
    
    def _save_l2_cache(self):
        try:
            with open(self.l2_cache_file, 'wb') as f:
                pickle.dump(self.l1_cache, f)
        except Exception as e:
            logging.warning(f"Cache save failed: {e}")
    
    def _compute_similarity(self, emb1: np.ndarray, emb2: np.ndarray) -> float:
        norm1 = np.linalg.norm(emb1)
        norm2 = np.linalg.norm(emb2)
        if norm1 == 0 or norm2 == 0:
            return 0.0
        return np.dot(emb1, emb2) / (norm1 * norm2)
    
    def get(self, query: str, query_embedding: Optional[np.ndarray] = None) -> Optional[Tuple[str, List[str]]]:
        query_hash = hashlib.md5(query.encode()).hexdigest()
        
        # Exact match
        if query_hash in self.l1_cache:
            entry = self.l1_cache[query_hash]
            self.l1_cache.move_to_end(query_hash)
            logging.info("Cache hit (exact)")
            return entry.response, entry.sources
        
        # Semantic similarity match
        if query_embedding is not None:
            for key, entry in self.l1_cache.items():
                similarity = self._compute_similarity(query_embedding, entry.embedding)
                if similarity >= CACHE_SIMILARITY_THRESHOLD:
                    logging.info(f"Cache hit (semantic: {similarity:.3f})")
                    return entry.response, entry.sources
        
        return None
    
    def set(self, query: str, query_embedding: np.ndarray, response: str, sources: List[str]):
        query_hash = hashlib.md5(query.encode()).hexdigest()
        
        if len(self.l1_cache) >= self.max_size:
            self.l1_cache.popitem(last=False)
        
        self.l1_cache[query_hash] = CacheEntry(
            query=query,
            embedding=query_embedding,
            response=response,
            sources=sources,
            timestamp=time.time()
        )
        
        if len(self.l1_cache) % 10 == 0:
            self._save_l2_cache()


class QueryClassifier:
    """Classify user queries for optimized retrieval"""
    
    FACTUAL_KEYWORDS = ["what", "who", "when", "where", "define", "explain", "list", "show", "tell"]
    ANALYTICAL_KEYWORDS = ["compare", "analyze", "why", "how can", "how does", "evaluate", "difference"]
    CONVERSATIONAL_KEYWORDS = ["hi", "hello", "thanks", "thank you", "bye", "hey"]
    
    @staticmethod
    def classify(query: str) -> Dict[str, any]:
        query_lower = query.lower().strip()
        word_count = len(query.split())
        
        is_pure_conversational = (
            word_count <= 3 and 
            any(query_lower.startswith(kw) for kw in QueryClassifier.CONVERSATIONAL_KEYWORDS)
        )
        
        intent = "factual"
        if is_pure_conversational:
            intent = "conversational"
        elif any(kw in query_lower for kw in QueryClassifier.ANALYTICAL_KEYWORDS):
            intent = "analytical"
        elif any(kw in query_lower for kw in QueryClassifier.FACTUAL_KEYWORDS):
            intent = "factual"
        
        complexity = min(1.0, (word_count / 20) + (0.3 if intent == "analytical" else 0))
        
        return {
            "intent": intent,
            "complexity": complexity,
            "word_count": word_count
        }


class EnhancedRAGSystem:
    """Enhanced RAG system with support for all file types"""
    
    def __init__(self):
        self.cache = MultiLevelCache()
        self.index = None
        self.initialized = False
        
        if RAG_AVAILABLE:
            try:
                self._initialize_settings()
                self.initialized = True
                logging.info("RAG System initialized")
            except Exception as e:
                logging.error(f"RAG initialization failed: {e}")
                self.initialized = False
    
    def _initialize_settings(self):
        """Initialize LlamaIndex settings"""
        Settings.embed_model = HuggingFaceEmbedding(model_name=EMBEDDING_MODEL)
        Settings.chunk_size = CHUNK_SIZE
        Settings.chunk_overlap = CHUNK_OVERLAP
        
        # Try Groq first, fallback to Ollama
        try:
            if GROQ_API_KEY and GROQ_API_KEY != "your_groq_api_key_here":
                Settings.llm = GroqLLM(
                    model="llama-3.3-70b-versatile",
                    temperature=0.3,
                    max_tokens=600
                )
                logging.info("Using Groq for RAG")
            else:
                raise ValueError("No Groq API key")
        except Exception:
            Settings.llm = OllamaLLM(model="gemma2:2b", temperature=0.3)
            logging.info("Using Ollama for RAG")
    
    def extract_text_from_pdf(self, path: str) -> str:
        """Extract text from PDF with error handling"""
        try:
            text_content = []
            with pdfplumber.open(path) as pdf:
                for page_num, page in enumerate(pdf.pages, 1):
                    page_text = page.extract_text()
                    if page_text:
                        text_content.append(f"[Page {page_num}]\n{page_text}")
            
            result = "\n\n".join(text_content)
            logging.info(f"PDF extracted: {len(result)} chars from {len(text_content)} pages")
            return result
        except Exception as e:
            logging.error(f"PDF extraction failed: {e}")
            return ""
    
    def extract_text_from_image(self, path: str) -> str:
        """Extract text from images using OCR with enhanced preprocessing"""
        try:
            # Try with PIL first
            img = Image.open(path)
            
            # Convert to RGB if needed
            if img.mode != 'RGB':
                img = img.convert('RGB')
            
            # Extract text using pytesseract
            text = pytesseract.image_to_string(img, lang='eng')
            
            # If pytesseract fails or returns empty, try OpenCV
            if not text.strip():
                img_cv = cv2.imread(path)
                if img_cv is not None:
                    # Preprocessing for better OCR
                    gray = cv2.cvtColor(img_cv, cv2.COLOR_BGR2GRAY)
                    # Apply thresholding
                    _, thresh = cv2.threshold(gray, 0, 255, cv2.THRESH_BINARY + cv2.THRESH_OTSU)
                    text = pytesseract.image_to_string(thresh)
            
            if text.strip():
                logging.info(f"Image OCR extracted: {len(text)} chars")
                return text
            else:
                logging.warning(f"No text found in image: {os.path.basename(path)}")
                return f"[Image file: {os.path.basename(path)} - No text detected]"
                
        except Exception as e:
            logging.error(f"Image OCR failed for {path}: {e}")
            return f"[Image file: {os.path.basename(path)} - OCR failed]"
    
    def extract_text_from_docx(self, path: str) -> str:
        """Extract text from DOCX files"""
        try:
            doc = docx.Document(path)
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            
            # Also extract text from tables
            tables_text = []
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text for cell in row.cells)
                    if row_text.strip():
                        tables_text.append(row_text)
            
            all_text = "\n".join(paragraphs)
            if tables_text:
                all_text += "\n\n[Tables]\n" + "\n".join(tables_text)
            
            logging.info(f"DOCX extracted: {len(all_text)} chars")
            return all_text
        except Exception as e:
            logging.error(f"DOCX extraction failed: {e}")
            return ""
    
    def extract_text_from_txt(self, path: str) -> str:
        """Extract text from TXT files with encoding detection"""
        encodings = ['utf-8', 'latin-1', 'cp1252', 'iso-8859-1']
        
        for encoding in encodings:
            try:
                with open(path, "r", encoding=encoding) as f:
                    text = f.read()
                    logging.info(f"TXT extracted: {len(text)} chars (encoding: {encoding})")
                    return text
            except (UnicodeDecodeError, Exception):
                continue
        
        logging.error(f"TXT extraction failed for all encodings")
        return ""
    
    def extract_text_from_csv(self, path: str) -> str:
        """Extract text from CSV files with better formatting"""
        try:
            rows = []
            with open(path, newline='', encoding='utf-8') as csvfile:
                reader = csv.reader(csvfile)
                header = next(reader, None)
                
                if header:
                    rows.append("Headers: " + " | ".join(header))
                    rows.append("-" * 50)
                
                for idx, row in enumerate(reader, 1):
                    rows.append(f"Row {idx}: " + " | ".join(row))
                    if idx >= 1000:  # Limit rows for very large CSVs
                        rows.append(f"[... truncated at {idx} rows ...]")
                        break
            
            result = "\n".join(rows)
            logging.info(f"CSV extracted: {len(result)} chars")
            return result
        except Exception as e:
            logging.error(f"CSV extraction failed: {e}")
            return ""
    
    def extract_text_from_json(self, path: str) -> str:
        """Extract text from JSON files with formatted output"""
        try:
            with open(path, "r", encoding="utf-8") as f:
                data = json.load(f)
            
            # Pretty print JSON
            formatted = json.dumps(data, indent=2, ensure_ascii=False)
            logging.info(f"JSON extracted: {len(formatted)} chars")
            return formatted
        except Exception as e:
            logging.error(f"JSON extraction failed: {e}")
            return ""
    
    def extract_text_from_xml(self, path: str) -> str:
        """Extract text from XML files"""
        try:
            tree = ET.parse(path)
            root = tree.getroot()
            
            def extract_from_element(element, level=0):
                texts = []
                indent = "  " * level
                
                # Add element tag and attributes
                if element.attrib:
                    attrs = ", ".join(f"{k}={v}" for k, v in element.attrib.items())
                    texts.append(f"{indent}<{element.tag} {attrs}>")
                else:
                    texts.append(f"{indent}<{element.tag}>")
                
                # Add element text
                if element.text and element.text.strip():
                    texts.append(f"{indent}  {element.text.strip()}")
                
                # Process children
                for child in element:
                    texts.extend(extract_from_element(child, level + 1))
                
                return texts
            
            all_text = "\n".join(extract_from_element(root))
            logging.info(f"XML extracted: {len(all_text)} chars")
            return all_text
        except Exception as e:
            logging.error(f"XML extraction failed: {e}")
            return ""
    
    def extract_text_from_html(self, path: str) -> str:
        """Extract text from HTML files"""
        try:
            with open(path, "r", encoding="utf-8") as f:
                html_content = f.read()
            
            # Remove script and style tags
            html_content = re.sub(r'<script[^>]*>.*?</script>', '', html_content, flags=re.DOTALL)
            html_content = re.sub(r'<style[^>]*>.*?</style>', '', html_content, flags=re.DOTALL)
            
            # Remove HTML tags
            text = re.sub(r'<[^>]+>', '', html_content)
            
            # Clean up whitespace
            text = re.sub(r'\s+', ' ', text).strip()
            
            logging.info(f"HTML extracted: {len(text)} chars")
            return text
        except Exception as e:
            logging.error(f"HTML extraction failed: {e}")
            return ""
    
    def extract_text_from_markdown(self, path: str) -> str:
        """Extract text from Markdown files"""
        try:
            with open(path, "r", encoding="utf-8") as f:
                text = f.read()
            
            logging.info(f"Markdown extracted: {len(text)} chars")
            return text
        except Exception as e:
            logging.error(f"Markdown extraction failed: {e}")
            return ""
    
    def extract_text(self, path: str) -> str:
        """Extract text from various file formats"""
        ext = os.path.splitext(path)[1].lower()
        
        extractors = {
            ".pdf": self.extract_text_from_pdf,
            ".jpg": self.extract_text_from_image,
            ".jpeg": self.extract_text_from_image,
            ".png": self.extract_text_from_image,
            ".bmp": self.extract_text_from_image,
            ".gif": self.extract_text_from_image,
            ".tiff": self.extract_text_from_image,
            ".docx": self.extract_text_from_docx,
            ".doc": self.extract_text_from_docx,
            ".txt": self.extract_text_from_txt,
            ".csv": self.extract_text_from_csv,
            ".json": self.extract_text_from_json,
            ".xml": self.extract_text_from_xml,
            ".html": self.extract_text_from_html,
            ".htm": self.extract_text_from_html,
            ".md": self.extract_text_from_markdown,
        }
        
        extractor = extractors.get(ext)
        if extractor:
            text = extractor(path)
            if text:
                return self.clean_text(text)
            else:
                logging.warning(f"No text extracted from {os.path.basename(path)}")
                return f"[File: {os.path.basename(path)} - No content extracted]"
        else:
            logging.warning(f"Unsupported file type: {ext}")
            return f"[Unsupported file type: {ext}]"
    
    def clean_text(self, text: str) -> str:
        """Clean and normalize text"""
        try:
            text = text.encode("latin1").decode("utf-8")
        except UnicodeError:
            pass

        text = re.sub(r'\s+', ' ', text)

        return text.strip()

    
    def create_index(self, file_paths: List[str]) -> bool:
        """Create vector index from documents"""
        if not self.initialized:
            logging.error("RAG System not initialized")
            return False
        
        try:
            documents = []
            for path in file_paths:
                if not os.path.exists(path):
                    logging.warning(f"File not found: {path}")
                    continue
                
                text = self.extract_text(path)
                if text and len(text.strip()) > 0:
                    doc = Document(
                        text=text,
                        metadata={
                            "source": os.path.basename(path),
                            "path": path,
                            "type": os.path.splitext(path)[1].lower()
                        }
                    )
                    documents.append(doc)
                    logging.info(f"Processed: {os.path.basename(path)} ({len(text)} chars)")
                else:
                    logging.warning(f"No content from: {os.path.basename(path)}")
            
            if not documents:
                logging.warning("No valid documents to index")
                return False
            
            logging.info(f"Creating index from {len(documents)} documents...")
            self.index = VectorStoreIndex.from_documents(documents, show_progress=True)
            logging.info("Index created successfully")
            return True
            
        except Exception as e:
            logging.error(f"Index creation failed: {traceback.format_exc()}")
            return False
    
    def compress_context(self, chunks: List[str], max_tokens: int = 1200) -> str:
        """Compress context by removing redundancy"""
        compressed = []
        token_count = 0
        
        for chunk in chunks:
            chunk_tokens = len(chunk.split())
            if token_count + chunk_tokens > max_tokens:
                remaining = max_tokens - token_count
                compressed.append(" ".join(chunk.split()[:remaining]))
                break
            compressed.append(chunk)
            token_count += chunk_tokens
        
        # Remove duplicate sentences
        unique_sentences = []
        seen = set()
        for chunk in compressed:
            try:
                for sent in sent_tokenize(chunk):
                    sent_clean = sent.strip().lower()
                    if sent_clean not in seen and len(sent_clean) > 20:
                        seen.add(sent_clean)
                        unique_sentences.append(sent)
            except:
                unique_sentences.append(chunk)
        
        return " ".join(unique_sentences)
    
    def query(self, query: str) -> Dict[str, any]:
        """Query the RAG system"""
        if not self.initialized or not self.index:
            return {
                "success": False,
                "response": "RAG system not initialized or no documents indexed",
                "sources": []
            }
        
        try:
            start_time = time.time()
            
            # Classify query
            classification = QueryClassifier.classify(query)
            
            # Check cache
            query_embedding = Settings.embed_model.get_text_embedding(query)
            cached = self.cache.get(query, np.array(query_embedding))
            
            if cached:
                response, sources = cached
                return {
                    "success": True,
                    "response": response,
                    "sources": sources,
                    "cached": True,
                    "time": time.time() - start_time
                }
            
            # Handle conversational
            if classification['intent'] == 'conversational':
                response = "Hello! I can help you with your documents. What would you like to know?"
                self.cache.set(query, np.array(query_embedding), response, [])
                return {
                    "success": True,
                    "response": response,
                    "sources": [],
                    "time": time.time() - start_time
                }
            
            # Retrieve documents
            k_retrieve = K_INIT if classification['complexity'] > 0.5 else max(K_FINAL, 8)
            retriever = self.index.as_retriever(similarity_top_k=k_retrieve)
            retrieved_nodes = retriever.retrieve(query)
            
            if not retrieved_nodes:
                return {
                    "success": True,
                    "response": "No relevant information found in the uploaded documents.",
                    "sources": [],
                    "time": time.time() - start_time
                }
            
            # Rerank if needed
            top_score = retrieved_nodes[0].score if hasattr(retrieved_nodes[0], 'score') else 1.0
            
            if top_score < SIMILARITY_THRESHOLD and len(retrieved_nodes) > K_FINAL:
                try:
                    reranker = CrossEncoder('cross-encoder/ms-marco-MiniLM-L-6-v2')
                    pairs = [[query, node.text] for node in retrieved_nodes[:K_RERANK]]
                    scores = reranker.predict(pairs)
                    
                    scored_nodes = list(zip(retrieved_nodes[:K_RERANK], scores))
                    scored_nodes.sort(key=lambda x: x[1], reverse=True)
                    final_nodes = [node for node, _ in scored_nodes[:K_FINAL]]
                except Exception:
                    final_nodes = retrieved_nodes[:K_FINAL]
            else:
                final_nodes = retrieved_nodes[:K_FINAL]
            
            # Compress context
            contexts = [node.text for node in final_nodes]
            compressed_context = self.compress_context(contexts)
            
            # Generate response
            prompt = (
                "You are a helpful assistant. Use ONLY the provided context to answer the question.\n"
                "If the information is not in the context, clearly state 'Information not available in the provided documents.'\n"
                "Be specific and cite relevant details from the context.\n\n"
                f"Context:\n{compressed_context}\n\n"
                f"Question: {query}\n\n"
                "Answer (be concise and specific):"
            )
            
            response_obj = Settings.llm.complete(prompt)
            response = response_obj.text.strip()
            
            # Extract sources
            sources = []
            for node in final_nodes[:3]:
                source = node.metadata.get('source', 'Unknown')
                file_type = node.metadata.get('type', '')
                sources.append(f"{source} ({file_type})")
            
            sources = list(dict.fromkeys(sources))
            
            # Cache result
            self.cache.set(query, np.array(query_embedding), response, sources)
            
            logging.info(f"Query completed in {time.time() - start_time:.2f}s")
            
            return {
                "success": True,
                "response": response,
                "sources": sources,
                "cached": False,
                "time": time.time() - start_time
            }
            
        except Exception as e:
            logging.error(f"RAG query failed: {traceback.format_exc()}")
            return {
                "success": False,
                "response": f"Error processing query: {str(e)}",
                "sources": []
            }


# Initialize enhanced RAG system
rag_system = EnhancedRAGSystem() if RAG_AVAILABLE else None

def detect_language(text: str) -> str:
    """Detect language: english, hindi, or hinglish"""
    cleaned_text = re.sub(r'[^\w\s]', '', text.lower())
    hindi_pattern = re.compile(r'[\u0900-\u097F]')
    
    hinglish_indicators = [
        'kya', 'hai', 'hoon', 'hain', 'aur', 'ki', 'ka', 'ke', 'ko', 'mein',
        'main', 'se', 'par', 'theek', 'nahi', 'haan', 'bahut', 'bohot'
    ]
    
    english_indicators = [
        'the', 'is', 'are', 'was', 'were', 'what', 'when', 'where', 'how',
        'feel', 'feeling', 'help', 'need'
    ]
    
    hindi_chars = len(hindi_pattern.findall(text))
    total_chars = len(re.sub(r'\s', '', text))
    
    if total_chars > 0 and (hindi_chars / total_chars) > 0.3:
        return 'hindi'
    
    words = cleaned_text.split()
    hinglish_count = sum(1 for word in words if word in hinglish_indicators)
    english_count = sum(1 for word in words if word in english_indicators)
    
    if hinglish_count > 0 and english_count > 0:
        return 'hinglish'
    elif hinglish_count > english_count:
        return 'hinglish'
    elif hindi_chars > 0:
        return 'hindi'
    return 'english'


def detect_emotion(text: str) -> Tuple[str, float]:
    """Detect emotion with confidence"""
    text_lower = text.lower()
    
    emotion_keywords = {
        "overwhelmed": ["overwhelmed", "burnout", "exhausted", "drained"],
        "sad": ["sad", "depressed", "lonely", "hopeless", "crying"],
        "angry": ["angry", "rage", "furious", "frustrated"],
        "anxious": ["anxious", "anxiety", "panic", "worried", "scared"],
        "happy": ["happy", "joy", "excited", "great"]
    }
    
    emotion_scores = {}
    for emotion, keywords in emotion_keywords.items():
        score = sum(len(kw.split()) for kw in keywords if kw in text_lower)
        emotion_scores[emotion] = score
    
    if max(emotion_scores.values()) == 0:
        return "neutral", 0.3
    
    detected = max(emotion_scores, key=emotion_scores.get)
    confidence = min(emotion_scores[detected] / 5.0, 1.0)
    
    return detected, confidence


def detect_high_risk(text: str) -> bool:
    """Detect crisis keywords"""
    crisis_keywords = [
        "suicidal", "kill myself", "want to die", "end it all", "suicide"
    ]
    text_lower = text.lower()
    return any(kw in text_lower for kw in crisis_keywords)


def call_ollama(prompt: str, model: str = "gemma2:2b") -> Optional[str]:
    """Call Ollama API"""
    try:
        response = requests.post(
            OLLAMA_API_URL,
            json={
                "model": model,
                "prompt": prompt,
                "stream": False,
                "options": {"temperature": 0.7, "top_p": 0.9, "num_predict": 2000}
            },
            timeout=60
        )
        
        if response.status_code == 200:
            return response.json().get("response", "").strip()
        return None
    except requests.exceptions.ConnectionError:
        logging.error("Ollama not running")
        return None
    except Exception as e:
        logging.error(f"Ollama error: {e}")
        return None

def allowed_file(filename: str) -> bool:
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS


def remove_file_safely(filepath: str):
    try:
        if filepath and os.path.exists(filepath):
            os.remove(filepath)
            logging.info(f"Removed: {os.path.basename(filepath)}")
    except Exception as e:
        logging.warning(f"Remove failed: {e}")


# Flask routes
@app.route('/')
def home():
    return render_template('home.html')


@app.route('/index')
def index():
    return render_template('index.html')

@app.route('/learn_more')
def learn_more():
    return render_template('learn_more.html')


@app.route('/disclaimer')
def disclaimer():
    return render_template('disclaimer.html')


@app.route('/about')
def about():
    return render_template('about.html')


@app.route('/resource')
def resource():
    return render_template('resource.html')


@app.route('/contact')
def contact():
    return render_template('contact.html')

@app.route('/help')
def help_page():
    return render_template('help.html')

@app.route('/upload', methods=['POST'])
def upload_file():
    """Enhanced file upload with better validation and feedback"""
    try:
        files = request.files.getlist('file')
        if not files:
            return jsonify({"success": False, "error": "No files provided"}), 400
        
        saved_files = []
        file_info = []
        
        for file in files:
            if not file or file.filename == '':
                continue
            
            if not allowed_file(file.filename):
                logging.warning(f"Rejected file: {file.filename} (unsupported type)")
                continue
            
            filename = secure_filename(file.filename)
            filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)
            
            # Handle duplicate filenames
            if os.path.exists(filepath):
                base, ext = os.path.splitext(filename)
                filename = f"{base}_{int(time.time())}{ext}"
                filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)
            
            file.save(filepath)
            file_size = os.path.getsize(filepath)
            saved_files.append(filepath)
            file_info.append({
                "name": filename,
                "path": filepath,
                "size": file_size,
                "type": os.path.splitext(filename)[1].lower()
            })
            logging.info(f"Uploaded: {filename} ({file_size} bytes)")
        
        if not saved_files:
            return jsonify({"success": False, "error": "No valid files uploaded"}), 400
        
        return jsonify({
            "success": True,
            "filepaths": saved_files,
            "filepath": saved_files[0],
            "count": len(saved_files),
            "files": file_info
        })
    
    except Exception as e:
        logging.error(f" Upload error: {traceback.format_exc()}")
        return jsonify({"success": False, "error": str(e)}), 500


@app.route('/search', methods=['POST'])
def search():
    """Enhanced search endpoint with improved RAG integration and conversation memory"""
    try:
        start_time = time.time()
        
        data = request.get_json(force=True)
        question = data.get('query') or data.get('question')
        raw_filepaths = data.get('filepaths') or data.get('filepath') or []
        session_name = data.get('session_name', 'default_session')
        
        if isinstance(raw_filepaths, str):
            raw_filepaths = [raw_filepaths]
        
        if not question:
            return jsonify({"success": False, "error": "Query required"}), 400
        
        logging.info(f"Query: {question[:100]} | Session: {session_name}")
        
        # Load conversation history
        conversation_file = os.path.join(CONVERSATION_FOLDER, f"{session_name}.json")
        conversation_history = []
        
        if os.path.exists(conversation_file):
            try:
                with open(conversation_file, 'r', encoding='utf-8') as f:
                    conversation_history = json.load(f)
                logging.info(f"Loaded {len(conversation_history)} previous messages")
            except Exception as e:
                logging.warning(f"Failed to load conversation history: {e}")
                conversation_history = []
        
        # Language & emotion detection
        detected_lang = detect_language(question)
        emotion, confidence = detect_emotion(question)
        
        emotion_emojis = {
            "overwhelmed": "😰", "sad": "😢", "angry": "😠",
            "anxious": "😨", "neutral": "😌", "happy": "😊"
        }
        emotion_emoji = emotion_emojis.get(emotion, "💭")
        
        logging.info(f"{detected_lang} |{emotion} {emotion_emoji}")
        
        # Crisis check
        if detect_high_risk(question):
            crisis_msg = {
                "english": "Crisis detected. You're not alone.\n\nIndia: +91 9152987821\nInternational: 988\n\nHelp is available 24/7.",
                "hindi": "संकट। आप अकेले नहीं।\n\nभारत: +91 9152987821\nअंतर्राष्ट्रीय: 988\n\nसहायता उपलब्ध है।",
                "hinglish": "Crisis.\n\nIndia: +91 9152987821\nInternational: 988\n\nMadad available hai."
            }
            return jsonify({
                "success": True,
                "emotion_detected": emotion,
                "language": detected_lang,
                "response": crisis_msg.get(detected_lang, crisis_msg["english"])
            })
        rag_context = ""
        rag_sources = []
        used_files = []
        
        if raw_filepaths and rag_system and rag_system.initialized:
            try:
                # Resolve file paths
                abs_paths = []
                for fp in raw_filepaths:
                    if os.path.isabs(fp):
                        abs_paths.append(fp)
                    else:
                        potential_path = os.path.join(app.config['UPLOAD_FOLDER'], os.path.basename(fp))
                        if os.path.exists(potential_path):
                            abs_paths.append(potential_path)
                
                # Filter existing files
                abs_paths = [fp for fp in abs_paths if os.path.exists(fp)]
                
                if abs_paths:
                    logging.info(f"Processing {len(abs_paths)} file(s) for RAG")
                    
                    # Create index
                    if rag_system.create_index(abs_paths):
                        # Query the RAG system
                        rag_result = rag_system.query(question)
                        
                        if rag_result["success"]:
                            rag_context = rag_result["response"]
                            rag_sources = rag_result.get("sources", [])
                            
                            cached_status = " (cached)" if rag_result.get("cached", False) else ""
                            logging.info(f"RAG response: {len(rag_context)} chars{cached_status}")
                        else:
                            logging.warning(f"RAG query failed: {rag_result.get('response')}")
                    else:
                        logging.warning("RAG index creation failed")
                    
                    used_files = abs_paths
                else:
                    logging.warning(" No valid file paths found")
            
            except Exception as e:
                logging.error(f"RAG processing error: {traceback.format_exc()}")
        if rag_context and not "not available" in rag_context.lower():
            if conversation_history:
                recent_history = conversation_history[-6:]  # Last 3 Q&A pairs
                history_context = "\n".join([
                    f"{'User' if msg['role'] == 'user' else 'Assistant'}: {msg['content']}"
                    for msg in recent_history
                ])
                
                # Add history context to RAG prompt
                enhanced_prompt = (
                    f"Previous conversation:\n{history_context}\n\n"
                    f"Current question: {question}\n\n"
                    f"Document information:\n{rag_context}\n\n"
                    "Answer the current question using the document information and conversation context."
                )
                
                try:
                    # Re-query with history context
                    if detected_lang == "english" and groq_client:
                        response = groq_client.chat.completions.create(
                            model="llama-3.3-70b-versatile",
                            messages=[
                                {"role": "system", "content": "You are a helpful assistant with access to documents and conversation history."},
                                {"role": "user", "content": enhanced_prompt}
                            ],
                            temperature=0.7,
                            max_tokens=2000
                        )
                        rag_context = response.choices[0].message.content.strip()
                except Exception as e:
                    logging.warning(f"History context enhancement failed: {e}")
            
            # Add sources to response
            if rag_sources:
                rag_context += "\n\nSources: " + ", ".join(rag_sources)
            
            # Save conversation
            conversation_history.append({"role": "user", "content": question})
            conversation_history.append({"role": "assistant", "content": rag_context})
            
            try:
                with open(conversation_file, 'w', encoding='utf-8') as f:
                    json.dump(conversation_history, f, indent=2, ensure_ascii=False)
            except Exception as e:
                logging.warning(f"Failed to save conversation: {e}")
            
            # Cleanup files
            for fp in used_files:
                remove_file_safely(fp)
            
            total_time = round(time.time() - start_time, 2)
            logging.info(f"Completed in {total_time}s (RAG only)")
            
            return jsonify({
                "success": True,
                "emotion_detected": emotion,
                "emotion_emoji": emotion_emoji,
                "emotion_confidence": round(confidence, 2),
                "language": detected_lang,
                "model_used": "rag-system",
                "rag_used": True,
                "sources": rag_sources,
                "response": f"{rag_context}\n\n{total_time}s"
            })
        system_prompts = {
            "english": prompt.prompts,
            "hindi": prompt.prompts,
            "hinglish": prompt.prompts
        }
        
        system_prompt = system_prompts.get(detected_lang, system_prompts["english"])
        user_message = question
        
        if rag_context and "not available" in rag_context.lower():
            user_message += "\n\n[Note: The uploaded documents don't contain relevant information for this question]"
        
        if emotion != "neutral" and confidence > 0.5:
            user_message += f"\n\n[Detected emotion: {emotion}, confidence: {confidence:.2f}]"
        assistant_reply = ""
        model_used = ""
        
        if detected_lang == "english":
            if not groq_client:
                return jsonify({
                    "success": False,
                    "error": "Groq unavailable",
                    "response": "English support requires Groq API."
                }), 500
            
            try:
                logging.info("Using Groq with conversation history")
                messages = [{"role": "system", "content": system_prompt}]
                if conversation_history:
                    recent_history = conversation_history[-10:]
                    for msg in recent_history:
                        messages.append({
                            "role": msg["role"],
                            "content": msg["content"]
                        })
                
                # Add current message
                messages.append({"role": "user", "content": user_message})
                
                response = groq_client.chat.completions.create(
                    model="llama-3.3-70b-versatile",
                    messages=messages,
                    temperature=0.7,
                    top_p=0.9,
                    max_tokens=2000
                )
                assistant_reply = response.choices[0].message.content.strip()
                model_used = "groq-llama-3.3-70b"
                logging.info(f"Groq response: {len(assistant_reply)} chars")
            
            except Exception as e:
                logging.error(f"Groq error: {e}")
                return jsonify({
                    "success": False,
                    "error": f"Groq error: {str(e)}",
                    "response": "Response generation failed."
                }), 500
        
        else:  # Hindi/Hinglish
            try:
                logging.info(f"Using Ollama for {detected_lang} with history")
                
                # Build prompt with conversation history
                history_text = ""
                if conversation_history:
                    recent_history = conversation_history[-6:]  # Last 3 exchanges
                    history_text = "Previous conversation:\n"
                    for msg in recent_history:
                        role_label = "User" if msg["role"] == "user" else "Assistant"
                        history_text += f"{role_label}: {msg['content']}\n"
                    history_text += "\n"
                
                ollama_prompt = f"{system_prompt}\n\n{history_text}Current question: {user_message}\n\nResponse:"
                assistant_reply = call_ollama(ollama_prompt, model="gemma2:2b")
                
                if assistant_reply is None:
                    return jsonify({
                        "success": False,
                        "error": "Ollama unavailable",
                        "response": "Hindi/Hinglish support requires Ollama.\nStart: ollama serve\nInstall model: ollama pull gemma2:2b"
                    }), 500
                
                model_used = "gemma2:2b"
                logging.info(f"Ollama response: {len(assistant_reply)} chars")
            
            except Exception as e:
                logging.error(f"Ollama error: {e}")
                return jsonify({
                    "success": False,
                    "error": f"Ollama error: {str(e)}",
                    "response": "Response generation failed."
                }), 500
        
        # Save conversation history
        conversation_history.append({"role": "user", "content": question})
        conversation_history.append({"role": "assistant", "content": assistant_reply})
        
        try:
            with open(conversation_file, 'w', encoding='utf-8') as f:
                json.dump(conversation_history, f, indent=2, ensure_ascii=False)
            logging.info(f"Saved conversation history ({len(conversation_history)} messages)")
        except Exception as e:
            logging.warning(f"Failed to save conversation: {e}")
        
        # Add RAG sources if available
        if rag_sources:
            assistant_reply += "\n\nReferenced: " + ", ".join(rag_sources)
        
        # Cleanup
        for fp in used_files:
            remove_file_safely(fp)
        
        total_time = round(time.time() - start_time, 2)
        logging.info(f"Completed in {total_time}s")
        
        return jsonify({
            "success": True,
            "emotion_detected": emotion,
            "emotion_emoji": emotion_emoji,
            "emotion_confidence": round(confidence, 2),
            "language": detected_lang,
            "model_used": model_used,
            "rag_used": bool(rag_context),
            "sources": rag_sources,
            "response": f"{assistant_reply}\n\n{total_time}s"
        })
    
    except Exception as e:
        logging.error(f"Search error: {traceback.format_exc()}")
        return jsonify({
            "success": False,
            "error": str(e),
            "response": "An error occurred processing your request."
        }), 500


@app.route('/send_email', methods=['POST'])
def send_email():
    try:
        if not (SENDER_EMAIL and SENDER_PASSWORD):
            raise ValueError("Email configuration missing")
        
        name = request.form["name"]
        email = request.form["email"]
        message = request.form["message"]
        
        msg = MIMEMultipart()
        msg["From"] = SENDER_EMAIL
        msg["To"] = RECEIVER_EMAIL
        msg["Subject"] = f"Contact Form: {name}"
        
        body = f"Name: {name}\nEmail: {email}\n\nMessage:\n{message}"
        msg.attach(MIMEText(body, "plain"))
        
        with smtplib.SMTP("smtp.gmail.com", 587) as server:
            server.starttls()
            server.login(SENDER_EMAIL, SENDER_PASSWORD)
            server.sendmail(SENDER_EMAIL, RECEIVER_EMAIL, msg.as_string())
        
        flash("Message sent successfully!", "success")
    except Exception as e:
        logging.error(f" Email error: {e}")
        flash(f"Error: {e}", "error")
    
    return redirect(url_for("contact"))


@app.route('/conversation/history', methods=['POST'])
def get_conversation_history():
    """Get conversation history for a session"""
    try:
        data = request.get_json(force=True)
        session_name = data.get('session_name', 'default_session')
        
        conversation_file = os.path.join(CONVERSATION_FOLDER, f"{session_name}.json")
        
        if os.path.exists(conversation_file):
            with open(conversation_file, 'r', encoding='utf-8') as f:
                history = json.load(f)
            return jsonify({
                "success": True,
                "history": history,
                "count": len(history)
            })
        else:
            return jsonify({
                "success": True,
                "history": [],
                "count": 0
            })
    except Exception as e:
        logging.error(f"Get history error: {e}")
        return jsonify({"success": False, "error": str(e)}), 500


@app.route('/conversation/clear', methods=['POST'])
def clear_conversation_history():
    """Clear conversation history for a session"""
    try:
        data = request.get_json(force=True)
        session_name = data.get('session_name', 'default_session')
        
        conversation_file = os.path.join(CONVERSATION_FOLDER, f"{session_name}.json")
        
        if os.path.exists(conversation_file):
            os.remove(conversation_file)
            logging.info(f"Cleared conversation: {session_name}")
        
        return jsonify({
            "success": True,
            "message": "Conversation history cleared"
        })
    except Exception as e:
        logging.error(f" Clear history error: {e}")
        return jsonify({"success": False, "error": str(e)}), 500


@app.route('/uploads/<path:filename>')
def uploaded_file(filename):
    return send_from_directory(app.config['UPLOAD_FOLDER'], filename)


@app.route('/health')
def health():
    """Health check endpoint"""
    ollama_ok = False
    try:
        resp = requests.get("http://localhost:11434/api/tags", timeout=2)
        ollama_ok = resp.status_code == 200
    except:
        pass
    
    return jsonify({
        "status": "online",
        "groq_available": groq_client is not None,
        "groq_installed": GROQ_AVAILABLE,
        "api_key_set": bool(GROQ_API_KEY and GROQ_API_KEY != "your_groq_api_key_here"),
        "ollama_available": ollama_ok,
        "rag_available": RAG_AVAILABLE,
        "rag_initialized": rag_system.initialized if rag_system else False,
        "supported_formats": list(ALLOWED_EXTENSIONS),
        "english_support": groq_client is not None,
        "hindi_hinglish_support": ollama_ok
    })

if __name__ == '__main__':
    print("CentrixSupport - Mental Health AI Chatbot")
    app.run(host="0.0.0.0", port=8000, debug=True, use_reloader=False)