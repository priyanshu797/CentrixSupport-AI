from __future__ import annotations

import csv
import hashlib
import json
import logging
import math
import os
import pickle
import re
import tempfile
import threading
import time
import zipfile
from collections import OrderedDict
from dataclasses import dataclass
from typing import ClassVar

import cv2
import docx
import nltk
import numpy as np
import pdfplumber
import pytesseract
import requests
import whisper
from groq import Groq as GroqClient
from llama_index.core import (
    Document,
    QueryBundle,
    Settings,
    StorageContext,
    VectorStoreIndex,
    load_index_from_storage,
)
from llama_index.embeddings.huggingface import HuggingFaceEmbedding
from nltk.tokenize import sent_tokenize
from sentence_transformers import CrossEncoder

BASE_DIR = os.path.dirname(os.path.abspath(__file__))

nltk.download("punkt", quiet=True)
nltk.download("punkt_tab", quiet=True)

# ── Module logger (fixes LOG015 — never use root logger directly) ──────────────
logger = logging.getLogger(__name__)

# ── Configuration ──────────────────────────────────────────────────────────────
PERSIST_DIR = os.path.join(os.path.dirname(BASE_DIR), "storage")
CACHE_DIR = os.path.join(BASE_DIR, "cache")
HASH_CACHE_FILE = os.path.join(os.path.dirname(BASE_DIR), "text_hash.txt")
EMBEDDING_MODEL = "sentence-transformers/all-MiniLM-L6-v2"
GROQ_MODEL = "llama-3.3-70b-versatile"

GROQ_API_KEY = "enter api key"

groq_client = GroqClient(api_key=GROQ_API_KEY)

CHUNK_SIZE = 512
CHUNK_OVERLAP = 50

# Retrieval parameters
K_INIT = 20
K_RERANK = 10
K_FINAL = 5
SIMILARITY_THRESHOLD = 0.3
DOC_RELEVANCE_THRESHOLD = 0.25
CACHE_SIZE = 1000
CACHE_SIMILARITY_THRESHOLD = 0.95

_reranker: CrossEncoder | None = None
_reranker_lock = threading.Lock()

logging.basicConfig(
    level=logging.WARNING,
    format="%(asctime)s - %(levelname)s - %(message)s",
)


# ── Cache System ───────────────────────────────────────────────────────────────
@dataclass
class CacheEntry:
    query: str
    embedding: np.ndarray
    response: str
    sources: list[str]
    timestamp: float
    scope: str = ""


class MultiLevelCache:
    """L1 (memory) + L2 (disk) caching with semantic similarity."""

    def __init__(self, cache_dir: str = CACHE_DIR, max_size: int = CACHE_SIZE) -> None:
        self.cache_dir = cache_dir
        self.max_size = max_size
        os.makedirs(cache_dir, exist_ok=True)
        self.l1_cache: OrderedDict[str, CacheEntry] = OrderedDict()
        self.l2_cache_file = os.path.join(cache_dir, "l2_cache.pkl")
        self._load_l2_cache()

    def _load_l2_cache(self) -> None:
        if os.path.exists(self.l2_cache_file):
            try:
                with open(self.l2_cache_file, "rb") as f:
                    self.l1_cache = pickle.load(f)
            except (OSError, pickle.UnpicklingError) as exc:
                logger.warning("Failed to load L2 cache: %s", exc)

    def _save_l2_cache(self) -> None:
        try:
            with open(self.l2_cache_file, "wb") as f:
                pickle.dump(self.l1_cache, f)
        except OSError as exc:
            logger.warning("Failed to save L2 cache: %s", exc)

    def _cosine(self, a: np.ndarray, b: np.ndarray) -> float:
        na, nb = float(np.linalg.norm(a)), float(np.linalg.norm(b))
        return float(np.dot(a, b) / (na * nb)) if na and nb else 0.0

    def get(
        self,
        query: str,
        query_embedding: np.ndarray | None = None,
        scope: str = "",
    ) -> tuple[str, list[str]] | None:
        key = hashlib.md5(f"{scope}\0{query}".encode()).hexdigest()
        if key in self.l1_cache:
            self.l1_cache.move_to_end(key)
            e = self.l1_cache[key]
            return e.response, e.sources
        if query_embedding is not None:
            for e in self.l1_cache.values():
                if (
                    getattr(e, "scope", "") == scope
                    and self._cosine(query_embedding, e.embedding)
                    >= CACHE_SIMILARITY_THRESHOLD
                ):
                    return e.response, e.sources
        return None

    def set(
        self,
        query: str,
        query_embedding: np.ndarray,
        response: str,
        sources: list[str],
        scope: str = "",
    ) -> None:
        key = hashlib.md5(f"{scope}\0{query}".encode()).hexdigest()
        if len(self.l1_cache) >= self.max_size:
            self.l1_cache.popitem(last=False)
        self.l1_cache[key] = CacheEntry(
            query=query,
            embedding=query_embedding,
            response=response,
            sources=sources,
            timestamp=time.time(),
            scope=scope,
        )
        if len(self.l1_cache) % 10 == 0:
            self._save_l2_cache()


cache = MultiLevelCache()


# ── Helpers ────────────────────────────────────────────────────────────────────
def initialize_settings() -> None:
    Settings.embed_model = HuggingFaceEmbedding(model_name=EMBEDDING_MODEL)
    Settings.chunk_size = CHUNK_SIZE
    Settings.chunk_overlap = CHUNK_OVERLAP
    logger.info("Embed model: %s | LLM: %s (via groq SDK)", EMBEDDING_MODEL, GROQ_MODEL)


def warm_reranker() -> bool:
    """Load the optional cross-encoder once so retrieval queries do not load it."""
    global _reranker
    if _reranker is not None:
        return True
    with _reranker_lock:
        if _reranker is None:
            try:
                _reranker = CrossEncoder("cross-encoder/ms-marco-MiniLM-L-6-v2")
                logger.info("Cross-encoder reranker ready")
            except Exception as exc:  # noqa: BLE001 — optional model warm-up
                logger.warning("Reranker warm-up failed: %s", exc)
                return False
    return True


def _groq_complete(
    prompt: str,
    temperature: float = 0.4,
    max_tokens: int = 700,
) -> str:
    """Single-turn completion through Groq SDK."""
    resp = groq_client.chat.completions.create(
        model=GROQ_MODEL,
        messages=[
            {
                "role": "system",
                "content": "You are a knowledgeable and helpful AI assistant. Answer questions clearly and accurately.",
            },
            {"role": "user", "content": prompt},
        ],
        temperature=temperature,
        max_completion_tokens=max_tokens,
    )
    return resp.choices[0].message.content or ""


# ── File Extractors ────────────────────────────────────────────────────────────
def extract_text_from_pdf(path: str) -> str:
    try:
        with pdfplumber.open(path) as pdf:
            return "\n".join([p.extract_text() or "" for p in pdf.pages])
    except OSError as exc:
        logger.error("PDF extraction failed: %s", exc)
        return ""


def extract_text_from_image(path: str) -> str:
    try:
        return pytesseract.image_to_string(cv2.imread(path))
    except OSError as exc:
        logger.error("Image OCR failed: %s", exc)
        return ""


def extract_text_from_docx(path: str) -> str:
    try:
        return "\n".join([p.text for p in docx.Document(path).paragraphs])
    except OSError as exc:
        logger.error("DOCX extraction failed: %s", exc)
        return ""


def extract_text_from_txt(path: str) -> str:
    for enc in ("utf-8", "latin-1", "cp1252"):
        try:
            with open(path, encoding=enc) as f:
                return f.read()
        except UnicodeDecodeError:
            continue
        except OSError as exc:
            logger.error("TXT extraction failed: %s", exc)
            return ""
    return ""


def extract_text_from_csv(path: str) -> str:
    try:
        with open(path, newline="", encoding="utf-8") as f:
            return "\n".join([" | ".join(row) for row in csv.reader(f)])
    except OSError as exc:
        logger.error("CSV extraction failed: %s", exc)
        return ""


def extract_text_from_json(path: str) -> str:
    try:
        with open(path, encoding="utf-8") as f:
            return json.dumps(json.load(f), indent=2)
    except OSError as exc:
        logger.error("JSON extraction failed: %s", exc)
        return ""


def extract_text_from_audio(path: str) -> str:
    try:
        return whisper.load_model("base").transcribe(path)["text"]
    except OSError as exc:
        logger.error("Audio transcription failed: %s", exc)
        return ""


def extract_text_from_zip(path: str) -> str:
    text = ""
    try:
        with zipfile.ZipFile(path, "r") as z, tempfile.TemporaryDirectory() as tmp:
            z.extractall(tmp)
            for root, _, files in os.walk(tmp):
                for name in files:
                    text += extract_text(os.path.join(root, name)) + "\n"
    except (OSError, zipfile.BadZipFile) as exc:
        logger.error("ZIP extraction failed: %s", exc)
    return text


def extract_text(path: str) -> str:
    ext = os.path.splitext(path)[1].lower()
    extractors = {
        ".pdf": extract_text_from_pdf,
        ".jpg": extract_text_from_image,
        ".jpeg": extract_text_from_image,
        ".png": extract_text_from_image,
        ".docx": extract_text_from_docx,
        ".txt": extract_text_from_txt,
        ".csv": extract_text_from_csv,
        ".json": extract_text_from_json,
        ".mp3": extract_text_from_audio,
        ".wav": extract_text_from_audio,
        ".mp4": extract_text_from_audio,
        ".zip": extract_text_from_zip,
    }
    fn = extractors.get(ext)
    return fn(path) if fn else ""


def clean_text(text: str) -> str:
    return re.sub(r"\s+", " ", text).replace("•", "-").replace("–", "-").strip()


def compute_hash(text: str) -> str:
    return hashlib.md5(text.encode("utf-8")).hexdigest()


def has_file_changed(text: str) -> bool:
    if os.path.exists(HASH_CACHE_FILE):
        with open(HASH_CACHE_FILE) as f:
            return compute_hash(text) != f.read().strip()
    return True


def cache_hash(text: str) -> None:
    with open(HASH_CACHE_FILE, "w") as f:
        f.write(compute_hash(text))


# ── Index Creation ─────────────────────────────────────────────────────────────
def create_index(file_paths: list[str]) -> VectorStoreIndex | None:
    docs = []
    for path in file_paths:
        text = clean_text(extract_text(path))
        if text:
            docs.append(
                Document(text=text, metadata={"source": os.path.basename(path)})
            )

    if not docs:
        logger.warning("No valid text extracted from files.")
        return None

    combined = "\n".join([d.text for d in docs])
    if os.path.exists(PERSIST_DIR) and not has_file_changed(combined):
        logger.info("Loading existing index…")
        ctx = StorageContext.from_defaults(persist_dir=PERSIST_DIR)
        loaded = load_index_from_storage(ctx)
        if not isinstance(loaded, VectorStoreIndex):
            logger.warning("Loaded index is not a VectorStoreIndex; rebuilding.")
        else:
            return loaded

    logger.info("Building new index…")
    index = VectorStoreIndex.from_documents(docs, show_progress=True)
    index.storage_context.persist(persist_dir=PERSIST_DIR)
    cache_hash(combined)
    return index


# ── Context Compression ────────────────────────────────────────────────────────
def compress_context(chunks: list[str], max_tokens: int = 1500) -> str:
    compressed, token_count = [], 0
    for chunk in chunks:
        n = len(chunk.split())
        if token_count + n > max_tokens:
            compressed.append(" ".join(chunk.split()[: max_tokens - token_count]))
            break
        compressed.append(chunk)
        token_count += n

    seen: set[str] = set()
    sentences: list[str] = []
    for chunk in compressed:
        for s in sent_tokenize(chunk):
            sc = s.strip().lower()
            if sc not in seen and len(sc) > 20:
                seen.add(sc)
                sentences.append(s)
    return " ".join(sentences)


# ── Tools ──────────────────────────────────────────────────────────────────────
class Tools:
    """Stateless tool collection used by the Agent."""

    @staticmethod
    def search_docs(
        index: VectorStoreIndex | list[VectorStoreIndex],
        query: str,
        query_embedding: list[float] | None = None,
    ) -> tuple[str | None, list[str], float]:
        """Hybrid retrieval: vector similarity + optional cross-encoder reranking."""
        indexes = index if isinstance(index, list) else [index]
        if query_embedding is None:
            query_embedding = Settings.embed_model.get_query_embedding(query)
        query_bundle = QueryBundle(query_str=query, embedding=query_embedding)
        nodes = []
        for document_index in indexes:
            retriever = document_index.as_retriever(similarity_top_k=K_INIT)
            nodes.extend(retriever.retrieve(query_bundle))
        nodes.sort(key=lambda node: float(node.score or 0.0), reverse=True)

        if not nodes:
            return None, [], 0.0

        best_score: float = (
            float(nodes[0].score)
            if hasattr(nodes[0], "score") and nodes[0].score is not None
            else 0.0
        )

        if best_score < SIMILARITY_THRESHOLD and len(nodes) > K_FINAL:
            try:
                if not warm_reranker() or _reranker is None:
                    raise RuntimeError("Cross-encoder reranker is unavailable")
                pairs = [[query, n.text] for n in nodes[:K_RERANK]]
                scores = _reranker.predict(pairs)
                ranked = sorted(
                    zip(nodes[:K_RERANK], scores),
                    key=lambda x: x[1],
                    reverse=True,
                )
                best_score = float(ranked[0][1]) if ranked else best_score
                nodes = [n for n, _ in ranked[:K_FINAL]]
            except (OSError, RuntimeError) as exc:
                logger.warning("Reranking failed: %s. Falling back to top-k.", exc)
                nodes = nodes[:K_FINAL]
        else:
            nodes = nodes[:K_FINAL]

        context = compress_context([n.text for n in nodes])
        sources = list(
            dict.fromkeys(n.metadata.get("source", "Unknown") for n in nodes[:3])
        )
        return context, sources, best_score

    @staticmethod
    def web_search(query: str) -> str:
        """DuckDuckGo instant-answer search (no API key required)."""
        try:
            headers = {"User-Agent": "Mozilla/5.0"}
            params: dict[str, str] = {
                "q": query,
                "format": "json",
                "no_html": "1",
                "skip_disambig": "1",
            }
            r = requests.get(
                "https://api.duckduckgo.com/",
                params=params,
                headers=headers,
                timeout=6,
            )
            data = r.json()
            results: list[str] = []
            if data.get("AbstractText"):
                results.append(data["AbstractText"])
            for topic in data.get("RelatedTopics", [])[:3]:
                if isinstance(topic, dict) and topic.get("Text"):
                    results.append(topic["Text"])
            if results:
                return "\n".join(results)
        except (OSError, requests.RequestException) as exc:
            logger.warning("Web search error: %s", exc)
        return ""

    @staticmethod
    def calculator(expr: str) -> str:
        try:
            safe = re.sub(r"[^0-9+\-*/().% ]", "", expr)
            result = eval(
                safe, {"__builtins__": {}}, {k: getattr(math, k) for k in dir(math)}
            )
            return str(result)
        except (SyntaxError, ArithmeticError, NameError, TypeError, ValueError) as exc:
            return f"Calculation error: {exc}"


# ── Query Classifier ───────────────────────────────────────────────────────────
class QueryClassifier:
    CONVERSATIONAL: ClassVar[list[str]] = [
        "hi",
        "hello",
        "hey",
        "thanks",
        "thank you",
        "bye",
        "good morning",
        "good night",
    ]
    ANALYTICAL: ClassVar[list[str]] = [
        "compare",
        "analyze",
        "why",
        "how can",
        "how does",
        "evaluate",
        "difference",
        "relationship",
    ]
    FACTUAL: ClassVar[list[str]] = [
        "what",
        "who",
        "when",
        "where",
        "define",
        "explain",
        "list",
        "show",
        "tell",
    ]
    TEMPORAL: ClassVar[list[str]] = [
        "latest",
        "current",
        "today",
        "news",
        "price",
        "live",
        "now",
        "2024",
        "2025",
        "2026",
    ]
    MATH_WORDS: ClassVar[list[str]] = [
        "calculate",
        "compute",
        "how many hours",
        "how many minutes",
        "how many days",
        "sum of",
        "total of",
    ]

    @classmethod
    def classify(cls, query: str) -> dict:
        q = query.lower().strip()
        words = q.split()
        wc = len(words)

        is_greeting = wc <= 4 and any(q.startswith(kw) for kw in cls.CONVERSATIONAL)
        has_number = bool(re.search(r"\d", q))
        has_math_op = any(op in q for op in ["+", "-", "*", "/", "^", "%"])
        is_math = (has_number and has_math_op) or (
            has_number and any(w in q for w in cls.MATH_WORDS)
        )
        is_temporal = any(kw in q for kw in cls.TEMPORAL)

        intent = "factual"
        if is_greeting:
            intent = "conversational"
        elif is_math:
            intent = "calculator"
        elif is_temporal:
            intent = "web_search"
        elif any(kw in q for kw in cls.ANALYTICAL):
            intent = "analytical"

        return {
            "intent": intent,
            "complexity": min(1.0, (wc / 20) + (0.3 if intent == "analytical" else 0)),
            "word_count": wc,
            "is_math": is_math,
            "is_temporal": is_temporal,
        }


# ── Agent ──────────────────────────────────────────────────────────────────────
class Agent:
    """
    Agentic RAG orchestrator.

    Decision flow per query
    ───────────────────────
    1. Cache hit?            → return cached answer
    2. Greeting?             → conversational reply
    3. Math expression?      → calculator tool
    4. Temporal / live data? → web search tool
    5. Has index?
       a. Retrieve from docs (hybrid vector + rerank)
       b. Score ≥ DOC_RELEVANCE_THRESHOLD?
          → build context-grounded answer
          → verify; if vague → enrich with web search
       c. Score too low (off-topic for the docs)?
          → answer from LLM general knowledge directly
          → if LLM is unsure → try web search as fallback
    6. No index?             → LLM general knowledge (+ web if needed)
    """

    def __init__(
        self,
        index: VectorStoreIndex | list[VectorStoreIndex] | None,
        cache_scope: str | None = None,
    ) -> None:
        self.index = index
        self.tools = Tools()
        self.cache_scope = cache_scope or (
            "documents" if index is not None else "general"
        )

    # ── Private helpers ────────────────────────────────────────────────────────
    def _llm(
        self, prompt: str, temperature: float = 0.4, max_tokens: int = 1200
    ) -> str:
        return _groq_complete(prompt, temperature=temperature, max_tokens=max_tokens)

    def _is_vague(self, text: str) -> bool:
        vague_phrases = [
            "i don't know",
            "not sure",
            "i cannot",
            "not available",
            "no information",
            "i'm unable",
            "cannot determine",
        ]
        return any(p in text.lower() for p in vague_phrases)

    def _extract_math_expr(self, query: str) -> str:
        prompt = (
            "Extract only the math expression from the query as a Python-evaluable string. "
            "Return ONLY the expression, no explanation.\n"
            f"Query: {query}\nExpression:"
        )
        expr = self._llm(prompt, temperature=0.0, max_tokens=60)
        return re.sub(r"[^0-9+\-*/().% ]", "", expr).strip()

    def _answer_from_docs(self, query: str, context: str) -> str:
        prompt = (
            "You are a helpful assistant. Use ONLY the provided document context to answer.\n"
            "Be specific, sufficiently detailed, and easy to understand. Do not omit "
            "relevant details merely to be brief. If the answer is not in the context, say exactly: "
            "'Not available in the documentation.'\n"
            "Formatting requirements:\n"
            "- Start with a direct answer in one or two sentences.\n"
            "- For multi-part answers, use descriptive Markdown headings (## Heading).\n"
            "- Put a blank line before and after every heading and paragraph.\n"
            "- Put every bullet or numbered item on its own line.\n"
            "- Explain important points in short paragraphs, not a dense wall of text.\n"
            "- End substantial answers with a '## Summary' section containing 2-4 concise bullets.\n"
            "- Do not output literal \\n characters; output real line breaks.\n\n"
            f"Context:\n{context}\n\nQuestion: {query}\nAnswer:"
        )
        return self._llm(prompt)

    def _answer_from_knowledge(self, query: str) -> str:
        prompt = (
            "You are a knowledgeable assistant. The user's question is not covered by any "
            "uploaded documents, so answer using your general knowledge.\n"
            "Be accurate, helpful, sufficiently detailed, and well summarized. Start with "
            "a direct answer. For multi-part answers, use descriptive Markdown headings, "
            "short paragraphs separated by blank lines, and one bullet per line. End a "
            "substantial answer with a '## Summary' section of 2-4 bullets. Use real line "
            "breaks and never return a dense wall of text.\n\n"
            f"Question: {query}\nAnswer:"
        )
        return self._llm(prompt, temperature=0.5)

    def _answer_with_web(self, query: str, web_context: str) -> str:
        prompt = (
            "Use the following web search results to answer the question accurately and in "
            "useful detail. Start with a direct answer. For multi-part answers, use Markdown "
            "headings, blank lines between short paragraphs, and one bullet per line. End "
            "substantial answers with a '## Summary' section of 2-4 bullets.\n\n"
            f"Web Results:\n{web_context}\n\nQuestion: {query}\nAnswer:"
        )
        return self._llm(prompt)

    # ── Public interface ───────────────────────────────────────────────────────
    def run(self, query: str) -> tuple[str, list[str], str]:
        """
        Returns (response, sources, tool_used).
        tool_used is one of: cache | conversational | calculator | web_search |
                              docs | docs+web | llm | llm+web
        """
        embedding = np.array(Settings.embed_model.get_text_embedding(query))
        cached = cache.get(query, embedding, scope=self.cache_scope)
        if cached:
            return cached[0], cached[1], "cache"

        classification = QueryClassifier.classify(query)

        # 2. Greeting / small-talk
        if classification["intent"] == "conversational":
            response = "Hello! How can I help you today? Feel free to ask me anything."
            cache.set(query, embedding, response, [], scope=self.cache_scope)
            return response, [], "conversational"

        # 3. Math
        if classification["intent"] == "calculator":
            expr = self._extract_math_expr(query)
            if expr:
                result = self.tools.calculator(expr)
                response = f"Result: {result}"
            else:
                response = "Could not extract a valid math expression from your query."
            cache.set(query, embedding, response, [], scope=self.cache_scope)
            return response, [], "calculator"

        # 4. Live / temporal data → web first
        if classification["intent"] == "web_search":
            web_ctx = self.tools.web_search(query)
            if web_ctx:
                response = self._answer_with_web(query, web_ctx)
                sources: list[str] = ["web"]
            else:
                response = self._answer_from_knowledge(query)
                sources = []
            cache.set(query, embedding, response, sources, scope=self.cache_scope)
            return response, sources, "web_search"

        # 5a. Document retrieval (if index exists)
        if self.index is not None:
            doc_context, doc_sources, best_score = self.tools.search_docs(
                self.index, query, embedding.tolist()
            )

            # 5b. Documents are relevant — ground answer in them
            if doc_context and best_score >= DOC_RELEVANCE_THRESHOLD:
                response = self._answer_from_docs(query, doc_context)

                if self._is_vague(response):
                    web_ctx = self.tools.web_search(query)
                    if web_ctx:
                        response = self._answer_with_web(query, web_ctx)
                        doc_sources.append("web (enriched)")
                    else:
                        response = self._answer_from_knowledge(query)

                cache.set(
                    query, embedding, response, doc_sources, scope=self.cache_scope
                )
                return response, doc_sources, "docs"

            # 5c. Off-topic → LLM general knowledge
            llm_response = self._answer_from_knowledge(query)
            if self._is_vague(llm_response):
                web_ctx = self.tools.web_search(query)
                if web_ctx:
                    llm_response = self._answer_with_web(query, web_ctx)
                    cache.set(
                        query, embedding, llm_response, ["web"], scope=self.cache_scope
                    )
                    return llm_response, ["web"], "llm+web"

            cache.set(query, embedding, llm_response, [], scope=self.cache_scope)
            return llm_response, [], "llm"

        # 6. No index → LLM general knowledge
        llm_response = self._answer_from_knowledge(query)
        fallback_sources: list[str] = []
        tool_used = "llm"

        if self._is_vague(llm_response):
            web_ctx = self.tools.web_search(query)
            if web_ctx:
                llm_response = self._answer_with_web(query, web_ctx)
                fallback_sources = ["web"]
                tool_used = "llm+web"

        cache.set(
            query,
            embedding,
            llm_response,
            fallback_sources,
            scope=self.cache_scope,
        )
        return llm_response, fallback_sources, tool_used


# ── Response Formatter ─────────────────────────────────────────────────────────
def format_response(response: str, sources: list[str]) -> str:
    response = response.strip()
    response = re.sub(
        r"^(Answer|Response|Here's|Based on):\s*", "", response, flags=re.IGNORECASE
    )
    response = re.sub(r"\*\*([^*]+)\*\*", r"\1", response)
    response = re.sub(r"\n{3,}", "\n\n", response)
    if sources:
        response += f"  [Source: {', '.join(dict.fromkeys(sources))}]"
    return response


# ── CLI Entry Point ────────────────────────────────────────────────────────────
def main() -> None:
    initialize_settings()

    file_input = input(
        "Enter file paths (comma-separated, or press Enter to skip): "
    ).strip()
    file_paths = [
        f.strip()
        for f in file_input.split(",")
        if f.strip() and os.path.exists(f.strip())
    ]

    index = None
    if file_paths:
        print(f"\nProcessing {len(file_paths)} file(s)…")
        index = create_index(file_paths)
        if index is None:
            print("Warning: index creation failed. Falling back to LLM-only mode.\n")
    else:
        print("No files provided — running in LLM-only mode.\n")

    agent = Agent(index)

    while True:
        query = input("Q: ").strip()
        if query.lower() in ("exit", "quit", "q"):
            print("Goodbye!")
            break
        if not query:
            continue

        t0 = time.time()
        response, sources, tool_used = agent.run(query)
        elapsed = (time.time() - t0) * 1000

        print(f"\n{format_response(response, sources)}")
        print(f"[Tool: {tool_used} | {elapsed:.0f}ms]\n")


if __name__ == "__main__":
    main()
